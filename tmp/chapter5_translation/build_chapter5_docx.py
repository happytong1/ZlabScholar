from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

root = Path(__file__).parent
template_path = root / "template.docx"
content_path = root / "chapter5_en.md"
figures_dir = root / "figures"
output_path = root / "AI+MBSE_Chapter5_English.docx"

expected_hash = "CBE927001D6C651D5B318C60B0BF4EEA29DFA012AB7AF3A9D773656B8CC9C353"
actual_hash = hashlib.sha256(template_path.read_bytes()).hexdigest().upper()
if actual_hash != expected_hash:
    raise RuntimeError(f"Template hash mismatch: {actual_hash}")

doc = Document(template_path)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_keep_together(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_together = value


def add_body(text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(28)
    p.paragraph_format.widow_control = True
    p.add_run(text)


def add_heading(text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if text == "References":
        p.paragraph_format.page_break_before = True
    p.paragraph_format.widow_control = True
    set_keep_with_next(p)
    p.add_run(text).bold = True


def image_dimensions(path: Path) -> tuple[int, int]:
    from PIL import Image

    with Image.open(path) as image:
        return image.size


def add_visual(filename: str, caption: str) -> None:
    path = figures_dir / filename
    if not path.exists():
        raise FileNotFoundError(path)
    px_w, px_h = image_dimensions(path)
    aspect = px_w / px_h
    max_w = 6.25
    max_h = 7.15
    width = min(max_w, max_h * aspect)
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    set_keep_with_next(p)
    set_keep_together(p)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))

    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = None
    c.paragraph_format.space_before = Pt(2)
    c.paragraph_format.space_after = Pt(8)
    c.paragraph_format.widow_control = True
    c.add_run(caption)


def add_reference(text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Pt(22)
    p.paragraph_format.first_line_indent = Pt(-22)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.widow_control = True
    for run in p.runs:
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)


blocks = [block.strip() for block in re.split(r"\n\s*\n", content_path.read_text(encoding="utf-8")) if block.strip()]
in_references = False
for block in blocks:
    if block.startswith("## "):
        heading = block[3:].strip()
        add_heading(heading)
        in_references = heading == "References"
    elif block.startswith("[FIGURE|") or block.startswith("[TABLE|"):
        kind, filename, caption = block.strip("[]").split("|", 2)
        add_visual(filename, caption)
    elif in_references and re.match(r"^\d+\.\s", block):
        add_reference(block)
    else:
        add_body(block.replace("\n", " "))

# Ask Word to refresh page-number fields on open while preserving the existing PAGE field.
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.save(output_path)
print(output_path)
