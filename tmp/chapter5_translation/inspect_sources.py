from pathlib import Path
from zipfile import ZipFile
from pypdf import PdfReader
from docx import Document

root = Path(__file__).parent
pdf_path = root / "source.pdf"
docx_path = root / "template.docx"

reader = PdfReader(str(pdf_path))
print(f"PDF_PAGES\t{len(reader.pages)}")
for index, page in enumerate(reader.pages, 1):
    text = page.extract_text() or ""
    compact = "".join(text.split())
    if any(key in compact for key in ("第五章", "第5章", "卫星敏捷设计")):
        print(f"PDF_MATCH\t{index}\t{compact[:240]}")

def walk_outlines(items, level=1):
    for item in items:
        if isinstance(item, list):
            walk_outlines(item, level + 1)
        else:
            try:
                page = reader.get_destination_page_number(item) + 1
            except Exception:
                page = -1
            print(f"TOC\t{level}\t{page}\t{getattr(item, 'title', str(item))}")

print("PDF_TOC")
try:
    walk_outlines(reader.outline)
except Exception as exc:
    print(f"TOC_ERROR\t{exc}")

doc = Document(str(docx_path))
print(f"DOCX_SECTIONS\t{len(doc.sections)}")
print(f"DOCX_PARAGRAPHS\t{len(doc.paragraphs)}")
print(f"DOCX_TABLES\t{len(doc.tables)}")
for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.replace("\t", "\\t").replace("\n", "\\n")
    if text or index < 20:
        print(f"P\t{index}\t{paragraph.style.name}\t{text}")
for table_index, table in enumerate(doc.tables):
    print(f"TABLE\t{table_index}\t{len(table.rows)}x{len(table.columns)}")
    for row_index, row in enumerate(table.rows):
        cells = [cell.text.replace("\t", "\\t").replace("\n", "\\n") for cell in row.cells]
        print(f"TR\t{table_index}\t{row_index}\t" + " || ".join(cells))

with ZipFile(docx_path) as package:
    media = [name for name in package.namelist() if name.startswith("word/media/")]
    headers = [name for name in package.namelist() if name.startswith("word/header")]
    footers = [name for name in package.namelist() if name.startswith("word/footer")]
    print(f"DOCX_MEDIA\t{len(media)}\t{','.join(media)}")
    print(f"DOCX_HEADERS\t{len(headers)}\t{','.join(headers)}")
    print(f"DOCX_FOOTERS\t{len(footers)}\t{','.join(footers)}")
