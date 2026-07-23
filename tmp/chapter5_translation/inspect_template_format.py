from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(__file__).parent / "template.docx"
doc = Document(path)

for i, p in enumerate(doc.paragraphs):
    pf = p.paragraph_format
    print(
        "P", i, repr(p.text), "style=", p.style.name,
        "align=", p.alignment,
        "before=", pf.space_before.pt if pf.space_before else None,
        "after=", pf.space_after.pt if pf.space_after else None,
        "line=", pf.line_spacing,
        "first=", pf.first_line_indent.pt if pf.first_line_indent else None,
        "page_before=", pf.page_break_before,
        "keep_next=", pf.keep_with_next,
    )
    for j, r in enumerate(p.runs):
        rpr = r._element.rPr
        east = None
        if rpr is not None and rpr.rFonts is not None:
            east = rpr.rFonts.get(qn("w:eastAsia"))
        print(
            "  R", j, repr(r.text), "font=", r.font.name, "east=", east,
            "size=", r.font.size.pt if r.font.size else None,
            "bold=", r.bold, "italic=", r.italic,
        )

for si, s in enumerate(doc.sections):
    print("SECTION", si, s.page_width.inches, s.page_height.inches, s.left_margin.inches, s.right_margin.inches, s.top_margin.inches, s.bottom_margin.inches)
    print("HEADER", [p.text for p in s.header.paragraphs])
    print("FOOTER", [p.text for p in s.footer.paragraphs])

for name in ("Normal", "Heading 1", "Heading 2", "Caption"):
    style = doc.styles[name]
    pf = style.paragraph_format
    print(
        "STYLE", name, "font=", style.font.name,
        "size=", style.font.size.pt if style.font.size else None,
        "bold=", style.font.bold,
        "align=", pf.alignment,
        "before=", pf.space_before.pt if pf.space_before else None,
        "after=", pf.space_after.pt if pf.space_after else None,
        "line=", pf.line_spacing,
        "first=", pf.first_line_indent.pt if pf.first_line_indent else None,
    )
